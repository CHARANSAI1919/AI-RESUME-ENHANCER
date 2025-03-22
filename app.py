from flask import Flask, request, jsonify, render_template, send_file
import requests
import spacy
import re
from docx import Document
from pdfminer.high_level import extract_text
import os
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer
from reportlab.lib.styles import getSampleStyleSheet

app = Flask(__name__, static_folder="static")

# Set your DeepSeek API key
DEEPSEEK_API_KEY = "sk-b7b8598094754056842f46825ad6f7b5"
DEEPSEEK_API_URL = "https://api.deepseek.com/v1/chat/completions"  # Updated with the correct endpoint

# Load SpaCy model
nlp = spacy.load("en_core_web_sm")

# Function to extract text and preserve formatting from a Word file
def extract_text_and_format_from_docx(file):
    doc = Document(file)
    full_text = ""
    paragraphs = []
    
    for para in doc.paragraphs:
        para_text = para.text.strip()
        if para_text:  # Only process non-empty paragraphs
            full_text += para_text + "\n"
            formatted_runs = []
            for run in para.runs:
                formatted_runs.append({
                    "text": run.text,
                    "bold": run.bold,
                    "italic": run.italic,
                    "underline": run.underline,
                    "font_size": run.font.size if run.font.size else None,
                    "font_name": run.font.name if run.font.name else None
                })
            
            paragraphs.append({
                "text": para_text,
                "style": para.style.name,  # Preserve paragraph style (e.g., Heading 1, Normal)
                "runs": formatted_runs  # Preserve text formatting
            })
    
    return full_text, paragraphs

# Function to reconstruct a Word document with corrected text
def reconstruct_docx(original_paragraphs, corrected_text, filename):
    doc = Document()
    
    # Split the corrected text into paragraphs
    corrected_paragraphs = corrected_text.split('\n')
    corrected_idx = 0
    
    # Reconstruct document with corrected text while preserving original formatting
    for para_info in original_paragraphs:
        if corrected_idx < len(corrected_paragraphs):
            # Create a new paragraph with the same style
            new_para = doc.add_paragraph(style=para_info["style"])
            
            # If the paragraph has multiple runs with different formatting
            if len(para_info["runs"]) > 1:
                corrected_para = corrected_paragraphs[corrected_idx]
                run_info = para_info["runs"][0]
                new_run = new_para.add_run(corrected_para)
                new_run.bold = run_info.get("bold", False)
                new_run.italic = run_info.get("italic", False)
                new_run.underline = run_info.get("underline", False)
                if run_info.get("font_name"):
                    new_run.font.name = run_info.get("font_name")
                if run_info.get("font_size"):
                    new_run.font.size = run_info.get("font_size")
            else:
                run_info = para_info["runs"][0] if para_info["runs"] else {}
                new_run = new_para.add_run(corrected_paragraphs[corrected_idx])
                new_run.bold = run_info.get("bold", False)
                new_run.italic = run_info.get("italic", False)
                new_run.underline = run_info.get("underline", False)
                if run_info.get("font_name"):
                    new_run.font.name = run_info.get("font_name")
                if run_info.get("font_size"):
                    new_run.font.size = run_info.get("font_size")
            
            corrected_idx += 1
    
    doc.save(filename)
    return filename

# Function to extract text from a PDF file
def extract_text_from_pdf(file_path):
    return extract_text(file_path)

# Function to save as PDF with better formatting
def save_as_pdf(text, filename):
    doc = SimpleDocTemplate(filename, pagesize=letter)
    styles = getSampleStyleSheet()
    
    story = []
    paragraphs = text.split('\n')
    
    for para in paragraphs:
        if para.strip():
            p = Paragraph(para, styles["Normal"])
            story.append(p)
            story.append(Spacer(1, 12))
    
    doc.build(story)
    return filename

# Function to correct grammar using DeepSeek API
def correct_grammar(text, job_description=""):
    headers = {
        "Authorization": f"Bearer {DEEPSEEK_API_KEY}",
        "Content-Type": "application/json",
    }
    
    prompt = """You are an expert resume editor and career coach. Enhance the resume text below by:
1. Fixing all grammar, spelling, and punctuation errors
2. Improving phrasing to sound more professional and impactful
3. Converting passive voice to active voice
4. Using stronger action verbs (e.g., change "responsible for" to "managed", "did" to "accomplished")
5. Quantifying achievements where possible
6. Keeping the same overall structure and content (don't add fictitious details)
7. Make sure all bullet points and sections remain intact

Respond ONLY with the improved resume text, with no explanations, comments, or additional formatting."""

    if job_description:
        prompt += "\n\nOptimize the resume to match this job description:\n" + job_description

    data = {
        "model": "deepseek-chat",  # Replace with the correct model name if needed
        "messages": [
            {"role": "system", "content": prompt},
            {"role": "user", "content": text}
        ],
        "max_tokens": 2000,
        "temperature": 0.2,  # Lower temperature for more conservative editing
    }
    
    try:
        response = requests.post(DEEPSEEK_API_URL, headers=headers, json=data)
        if response.status_code == 200:
            result = response.json()
            if "choices" in result and len(result["choices"]) > 0:
                return result["choices"][0]["message"]["content"].strip()
            else:
                return text  # Fallback to original text if API response format unexpected
        else:
            print(f"API Error: {response.status_code} - {response.text}")
            return text  # Fallback to original text if API fails
    except Exception as e:
        print(f"Exception when calling DeepSeek API: {e}")
        return text  # Fallback to original text if exception occurs

# Function to extract keywords using SpaCy
def extract_keywords(text):
    doc = nlp(text.lower())
    
    # Get both individual tokens and noun chunks (for multi-word terms)
    keywords = set()
    
    # Add relevant noun chunks (typically job titles, skills, etc.)
    for chunk in doc.noun_chunks:
        if not all(token.is_stop for token in chunk):
            keywords.add(chunk.text.strip())
    
    # Add individual relevant tokens
    for token in doc:
        if token.is_alpha and not token.is_stop and len(token.text) > 1:
            if token.pos_ in ["NOUN", "PROPN", "ADJ", "VERB"]:
                keywords.add(token.text.strip())
    
    return list(keywords)

# Function to suggest missing keywords
def suggest_keywords(resume_text, job_description):
    if not job_description.strip():
        return []
        
    resume_keywords = set(extract_keywords(resume_text))
    job_keywords = set(extract_keywords(job_description))
    
    # Focus on important keywords from job description
    important_missing = job_keywords - resume_keywords
    
    # Sort by importance (frequency in job description)
    job_desc_lower = job_description.lower()
    sorted_missing = sorted(important_missing, 
                           key=lambda x: job_desc_lower.count(x), 
                           reverse=True)
    
    # Return top 15 most important missing keywords
    return sorted_missing[:15]

# Function to count grammar issues
def count_grammar_issues(text):
    # Simple grammar checks (can be improved with NLP libraries)
    issues = 0
    
    # Check for common errors (simplified)
    common_errors = [
        r'\b(its|it\'s)\b',           # its/it's confusion
        r'\b(their|there|they\'re)\b', # their/there/they're confusion
        r'\b(your|you\'re)\b',         # your/you're confusion
        r'\b(affect|effect)\b',       # affect/effect confusion
        r'\b(then|than)\b',           # then/than confusion
        r'\ba\s+[aeiou]',             # "a" before vowel sound
        r'\ban\s+[^aeiou]',           # "an" before consonant sound
        r'\s\s+',                     # Multiple spaces
        r'[.!?][A-Za-z]',             # Missing space after punctuation
        r'[,;:][A-Za-z]',             # Missing space after comma
    ]
    
    for pattern in common_errors:
        issues += len(re.findall(pattern, text))
    
    # Check for sentence fragments and run-on sentences (simplified)
    sentences = re.split(r'[.!?]+', text)
    for sentence in sentences:
        if len(sentence.split()) < 3 and sentence.strip():  # Too short
            issues += 1
        elif len(sentence.split()) > 30 and sentence.strip():  # Too long
            issues += 1
    
    return min(issues, 20)  # Cap at 20 issues

# Function to check formatting quality
def check_formatting(paragraphs):
    formatting_issues = 0
    
    # Check formatting consistency
    styles_used = set()
    fonts_used = set()
    sizes_used = set()
    
    for para in paragraphs:
        styles_used.add(para["style"])
        for run in para["runs"]:
            if run.get("font_name"):
                fonts_used.add(run.get("font_name"))
            if run.get("font_size"):
                sizes_used.add(run.get("font_size"))
    
    # Too many different styles/fonts is a formatting issue
    if len(styles_used) > 5:
        formatting_issues += min(len(styles_used) - 5, 5)
    if len(fonts_used) > 2:
        formatting_issues += min(len(fonts_used) - 2, 5)
    if len(sizes_used) > 4:
        formatting_issues += min(len(sizes_used) - 4, 5)
    
    # Check for proper heading structure
    has_header = any("Heading" in style for style in styles_used)
    if not has_header:
        formatting_issues += 5
    
    # Too many or too few paragraphs
    if len(paragraphs) < 5:
        formatting_issues += 3
    elif len(paragraphs) > 50:
        formatting_issues += 3
    
    # Calculate score out of 20
    formatting_score = max(20 - formatting_issues, 0)
    return formatting_score

# Function to check resume structure
def check_structure(text):
    structure_score = 10  # Start with perfect score
    
    # Check for essential resume sections
    essential_sections = [
        r'\b(profile|summary|objective)\b',
        r'\b(experience|work|employment)\b',
        r'\b(education|academic|university|college)\b',
        r'\b(skills|competencies|expertise)\b',
    ]
    
    sections_found = 0
    for pattern in essential_sections:
        if re.search(pattern, text, re.IGNORECASE):
            sections_found += 1
    
    # Deduct points for missing sections
    if sections_found < len(essential_sections):
        structure_score -= (len(essential_sections) - sections_found) * 2
    
    # Check for contact information
    has_contact = re.search(r'\b(email|phone|tel|contact|\d{3}[-.]?\d{3}[-.]?\d{4})\b', text, re.IGNORECASE)
    if not has_contact:
        structure_score -= 2
    
    # Check for bullet points (good structure practice)
    bullets = len(re.findall(r'•|\*', text))
    if bullets < 5:
        structure_score -= 1
    
    return max(structure_score, 0)  # Don't go below 0

# Function to calculate action verbs usage
def check_action_verbs(text):
    text = text.lower()
    
    # Common strong action verbs in resumes
    strong_verbs = [
        'achieved', 'acquired', 'adapted', 'addressed', 'administered', 'advised', 'analyzed', 'applied',
        'appointed', 'approved', 'arranged', 'assembled', 'assessed', 'assigned', 'assisted', 'attained',
        'authored', 'balanced', 'budgeted', 'built', 'calculated', 'chaired', 'coached', 'collaborated',
        'communicated', 'compiled', 'completed', 'composed', 'computed', 'conducted', 'consolidated',
        'constructed', 'contracted', 'converted', 'coordinated', 'created', 'decreased', 'defined',
        'delegated', 'delivered', 'demonstrated', 'designed', 'developed', 'devised', 'directed',
        'distributed', 'documented', 'doubled', 'earned', 'edited', 'effected', 'eliminated', 'enlarged',
        'established', 'evaluated', 'examined', 'executed', 'expanded', 'expedited', 'facilitated',
        'finalized', 'formulated', 'founded', 'generated', 'guided', 'headed', 'identified', 'implemented',
        'improved', 'increased', 'initiated', 'inspected', 'installed', 'instituted', 'instructed',
        'interpreted', 'introduced', 'invented', 'launched', 'lectured', 'led', 'maintained', 'managed',
        'marketed', 'mediated', 'modernized', 'monitored', 'motivated', 'negotiated', 'operated',
        'organized', 'originated', 'overhauled', 'oversaw', 'performed', 'pioneered', 'planned',
        'prepared', 'presented', 'prioritized', 'processed', 'produced', 'programmed', 'proposed',
        'provided', 'published', 'purchased', 'recommended', 'reduced', 'reengineered', 'reorganized',
        'researched', 'resolved', 'revamped', 'reviewed', 'revised', 'scheduled', 'set up', 'simplified',
        'solved', 'spearheaded', 'standardized', 'streamlined', 'strengthened', 'structured', 'supervised',
        'supported', 'surpassed', 'trained', 'transformed', 'translated', 'trimmed', 'upgraded'
    ]
    
    # Weak verbs to avoid
    weak_verbs = [
        'was', 'were', 'am', 'is', 'are', 'been', 'did', 'made', 'got', 'used', 'worked', 'helped',
        'responsible for', 'duties included', 'handled', 'served as', 'involved in'
    ]
    
    strong_count = 0
    weak_count = 0
    
    for verb in strong_verbs:
        strong_count += len(re.findall(r'\b' + verb + r'\b', text))
    
    for verb in weak_verbs:
        weak_count += len(re.findall(r'\b' + verb + r'\b', text))
    
    total_verbs = strong_count + weak_count
    if total_verbs == 0:
        return []  # No verbs found
    
    weak_verb_suggestions = []
    if weak_count > 0:
        # Find the specific weak verbs used
        for verb in weak_verbs:
            if re.search(r'\b' + verb + r'\b', text):
                # Suggest replacement
                if verb == 'responsible for':
                    weak_verb_suggestions.append(f"Replace 'responsible for' with 'managed', 'led', or 'oversaw'")
                elif verb == 'duties included':
                    weak_verb_suggestions.append(f"Replace 'duties included' with specific actions like 'implemented', 'developed', or 'coordinated'")
                elif verb == 'worked':
                    weak_verb_suggestions.append(f"Replace 'worked' with more specific verbs like 'executed', 'performed', or 'conducted'")
                elif verb == 'helped':
                    weak_verb_suggestions.append(f"Replace 'helped' with 'assisted', 'supported', or 'facilitated'")
                elif verb in ['was', 'were', 'am', 'is', 'are', 'been']:
                    weak_verb_suggestions.append(f"Replace '{verb}' with action verbs that describe what you actually did")
                else:
                    weak_verb_suggestions.append(f"Replace '{verb}' with stronger action verbs")
    
    return weak_verb_suggestions

# Function to calculate resume score
def calculate_resume_score(resume_text, job_description, original_paragraphs):
    # Grammar Score (out of 40)
    grammar_issues = count_grammar_issues(resume_text)
    grammar_score = max(40 - grammar_issues * 2, 0)
    
    # Keyword matching Score (out of 30)
    missing_keywords = suggest_keywords(resume_text, job_description)
    
    # Calculate percentage of matching keywords
    if job_description.strip():
        job_keywords = extract_keywords(job_description)
        resume_keywords = extract_keywords(resume_text)
        
        if job_keywords:
            matching_keywords = set(resume_keywords).intersection(set(job_keywords))
            keyword_match_percentage = len(matching_keywords) / len(job_keywords)
            keyword_score = min(int(keyword_match_percentage * 30), 30)
        else:
            keyword_score = 15  # Default if no job description or no keywords found
    else:
        keyword_score = 15  # Default if no job description
    
    # Formatting Score (out of 20)
    formatting_score = check_formatting(original_paragraphs)
    
    # Structure Score (out of 10)
    structure_score = check_structure(resume_text)
    
    # Total Score
    total_score = grammar_score + keyword_score + formatting_score + structure_score
    
    # Feedback based on score
    if total_score >= 90:
        feedback = "Your resume is **Excellent**! It is well-optimized for the job description."
    elif total_score >= 80:
        feedback = "Your resume is **Very Good**. Some minor improvements can make it even better."
    elif total_score >= 70:
        feedback = "Your resume is **Good**. Several improvements are recommended."
    elif total_score >= 60:
        feedback = "Your resume is **Fair**. Significant improvements are needed."
    else:
        feedback = "Your resume needs **Major Improvements**. Consider a substantial revision."
    
    # Detailed score breakdown
    score_breakdown = {
        "grammar_score": grammar_score,
        "keyword_score": keyword_score,
        "formatting_score": formatting_score,
        "structure_score": structure_score,
        "total_score": total_score,
        "feedback": feedback
    }
    
    return score_breakdown

# Route for the homepage
@app.route("/")
def home():
    return render_template("index.html")

# Route to handle file upload and processing
@app.route("/enhance", methods=["POST"])
def enhance():
    # Get uploaded files
    resume_file = request.files["resume"]
    job_description = request.form.get("job_description", "")

    # Save the uploaded file temporarily
    temp_file_path = None
    try:
        # Create temp directory if it doesn't exist
        if not os.path.exists("temp"):
            os.makedirs("temp")
            
        if resume_file.filename.endswith(".docx"):
            # Save the file temporarily
            temp_file_path = os.path.join("temp", resume_file.filename)
            resume_file.save(temp_file_path)
            
            # Extract text and formatting from the Word file
            resume_text, paragraphs = extract_text_and_format_from_docx(resume_file)

            # Correct grammar and enhance the text
            corrected_text = correct_grammar(resume_text, job_description)

            # Reconstruct the Word document with corrected text
            corrected_file_path_docx = "corrected_resume.docx"
            reconstruct_docx(paragraphs, corrected_text, corrected_file_path_docx)

            # Save corrected text to a TXT file with UTF-8 encoding
            corrected_file_path_txt = "corrected_resume.txt"
            with open(corrected_file_path_txt, "w", encoding="utf-8") as f:
                f.write(corrected_text)

            # Prepare download links
            download_links = {
                "txt": f"/download/{corrected_file_path_txt}",
                "docx": f"/download/{corrected_file_path_docx}"
            }

        elif resume_file.filename.endswith(".pdf"):
            # Save the file temporarily
            temp_file_path = os.path.join("temp", resume_file.filename)
            resume_file.save(temp_file_path)
            
            # Extract text from the saved file
            resume_text = extract_text_from_pdf(temp_file_path)
            
            # Create empty paragraphs structure for PDF (since we can't extract formatting)
            paragraphs = [{"text": p, "style": "Normal", "runs": [{"text": p, "bold": False, "italic": False}]} 
                         for p in resume_text.split('\n') if p.strip()]

            # Correct grammar and enhance the text
            corrected_text = correct_grammar(resume_text, job_description)

            # Save corrected text to a PDF file
            corrected_file_path_pdf = "corrected_resume.pdf"
            save_as_pdf(corrected_text, corrected_file_path_pdf)

            # Save corrected text to a TXT file with UTF-8 encoding
            corrected_file_path_txt = "corrected_resume.txt"
            with open(corrected_file_path_txt, "w", encoding="utf-8") as f:
                f.write(corrected_text)

            # Prepare download links
            download_links = {
                "txt": f"/download/{corrected_file_path_txt}",
                "pdf": f"/download/{corrected_file_path_pdf}"
            }

        else:
            return jsonify({"error": "Unsupported file format. Please upload a PDF or Word file."})

        # Suggest missing keywords
        missing_keywords = suggest_keywords(corrected_text, job_description)

        # Calculate resume score with more nuanced metrics
        score_breakdown = calculate_resume_score(corrected_text, job_description, paragraphs)

        # Suggest action verbs
        action_verb_suggestions = check_action_verbs(resume_text)

        # Compare before and after
        diff = {
            "original_length": len(resume_text),
            "enhanced_length": len(corrected_text),
            "difference_percentage": round((len(corrected_text) - len(resume_text)) / len(resume_text) * 100, 1) if len(resume_text) > 0 else 0
        }

        # Return results
        return jsonify({
            "corrected_text": corrected_text,
            "original_text": resume_text,
            "missing_keywords": missing_keywords,
            "score_breakdown": score_breakdown,
            "action_verb_suggestions": action_verb_suggestions,
            "download_links": download_links,
            "diff_info": diff
        })
    except Exception as e:
        # Log the error for debugging
        import traceback
        print(f"Error: {str(e)}")
        print(traceback.format_exc())
        return jsonify({"error": f"An error occurred: {str(e)}"})
    finally:
        # Clean up temporary file
        if temp_file_path and os.path.exists(temp_file_path):
            os.remove(temp_file_path)

# Route to download corrected resume
@app.route("/download/<filename>")
def download(filename):
    return send_file(filename, as_attachment=True)

if __name__ == "__main__":
    # Create a temporary directory if it doesn't exist
    if not os.path.exists("temp"):
        os.makedirs("temp")
    app.run(debug=True)

# Changes Made in the Enhanced Resume:
# 1. **Grammar and Spelling Corrections**: Fixed all grammar, spelling, and punctuation errors.
# 2. **Improved Phrasing**: Enhanced the phrasing to sound more professional and impactful.
# 3. **Active Voice Usage**: Converted passive voice constructions to active voice.
# 4. **Stronger Action Verbs**: Replaced weak verbs with stronger action verbs to convey accomplishments more effectively.
# 5. **Quantified Achievements**: Added quantifiable results where possible to highlight achievements.
# 6. **Keyword Optimization**: Incorporated important keywords from the job description to improve relevance.
# 7. **Formatting Improvements**: Ensured consistent formatting throughout the document.
# 8. **Structure and Organization**: Improved the overall structure and organization of the resume for better readability.